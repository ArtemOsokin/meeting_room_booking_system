from datetime import datetime

from core.settings import REPORT_EXT, REPORT_FILE_PATH
from django.db.models import QuerySet
from docx import Document
from rest_framework.request import Request

import mrbs_app.api.exceptions as booking_exceptions
import mrbs_app.serializers.booking as booking_serializers
from mrbs_app.models import Reservation


class BookingService:
    @staticmethod
    def _get_active_reservations(
        room_id: int, reserved_from: datetime, reserved_to: datetime
    ) -> QuerySet[Reservation]:
        return Reservation.objects.filter(
            room_id=room_id,
            reserved_from__lt=reserved_to,
            reserved_to__gt=reserved_from,
            status=Reservation.ReservationStatus.ACTIVE,
        )

    @staticmethod
    def _check_reserved_time(reserved_from: datetime, reserved_to: datetime):
        if reserved_from >= reserved_to:
            raise booking_exceptions.ReservationTimeError

    def _is_room_available(
        self, room_id: int, reserved_from: datetime, reserved_to: datetime
    ) -> bool:
        return not self._get_active_reservations(
            room_id=room_id, reserved_from=reserved_from, reserved_to=reserved_to
        ).exists()

    def make_reservation(self, request: Request):
        reservation_serializer = booking_serializers.ReservationCreateSerializer(data=request.data)
        reservation_serializer.is_valid(raise_exception=True)
        data = reservation_serializer.data

        self._check_reserved_time(
            reserved_from=data['reserved_from'], reserved_to=data['reserved_to']
        )
        if self._is_room_available(
            room_id=data['room_id'],
            reserved_from=data['reserved_from'],
            reserved_to=data['reserved_to'],
        ):
            Reservation(
                purpose_of_booking=data['purpose_of_booking'],
                reserved_from=data['reserved_from'],
                reserved_to=data['reserved_to'],
                status=Reservation.ReservationStatus.ACTIVE,
                user=request.user,
                room_id=data['room_id'],
            ).save()
            return
        raise booking_exceptions.ReservationBusyError

    def get_reservations_by_period(self, request: Request) -> list[Reservation]:
        reservation_serializer = booking_serializers.ReservationsRequestSerializer(
            data=request.query_params.dict()
        )
        reservation_serializer.is_valid(raise_exception=True)
        data = reservation_serializer.data
        return self._get_active_reservations(
            room_id=data['room_id'],
            reserved_from=data['reserved_from'],
            reserved_to=data['reserved_to'],
        )


class BookingReportService:
    @staticmethod
    def _get_reservations(
        reserved_from: datetime,
        reserved_to: datetime,
        room_id: int = None,
    ) -> QuerySet[Reservation]:
        if room_id:
            return Reservation.objects.filter(
                reserved_from__lt=reserved_to, reserved_to__gt=reserved_from, room_id=room_id
            )
        return Reservation.objects.filter(
            reserved_from__lt=reserved_to,
            reserved_to__gt=reserved_from,
        )

    @staticmethod
    def _make_document(bookings: QuerySet[Reservation]) -> Document:
        doc = Document()
        doc.add_heading('Отчет о бронированиях переговорных комнат', level=0)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Кто бронировал'
        hdr_cells[2].text = 'Время бронирования'
        hdr_cells[3].text = 'Цель бронирования'
        hdr_cells[4].text = 'Номер комнаты'

        for booking in bookings:
            row_cells = table.add_row().cells
            row_cells[0].text = str(booking.id)
            row_cells[1].text = str(booking.user)
            row_cells[2].text = (
                f'{booking.reserved_from.strftime("%Y-%m-%d %H:%M")} - '
                f'{booking.reserved_to.strftime("%Y-%m-%d %H:%M")}'
            )
            row_cells[3].text = booking.purpose_of_booking
            row_cells[4].text = str(booking.room.number)
        doc.add_page_break()
        return doc

    def get_reservations_report(self, request: Request) -> str:
        reservation_serializer = booking_serializers.ReservReportRequestSerializer(
            data=request.query_params.dict()
        )
        reservation_serializer.is_valid(raise_exception=True)
        data = reservation_serializer.data

        bookings = self._get_reservations(
            room_id=data.get('room_id'),
            reserved_from=data['reserved_from'],
            reserved_to=data['reserved_to'],
        )

        document = self._make_document(bookings=bookings)

        file_path = (
            f'{REPORT_FILE_PATH}'
            f'{data["reserved_from"]}_'
            f'{data["reserved_to"]}'
            f'{REPORT_EXT}'
        )
        document.save(file_path)

        return file_path
